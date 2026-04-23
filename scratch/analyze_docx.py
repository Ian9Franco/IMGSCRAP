import sys
from docx import Document

# Forzar UTF-8 en la salida de consola
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

def analyze_docx(file_path):
    doc = Document(file_path)
    print(f"--- Análisis de {file_path} ---")
    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        if text:
            print(f"[{style}] {text}")
    
    for table in doc.tables:
        print("\n[Table found]")
        for row in table.rows:
            print(" | ".join(cell.text.strip() for cell in row.cells))

if __name__ == "__main__":
    analyze_docx(r"d:\Dev\CodeProjects\imgscrap\documentacion\69-Necochea 437, Piso 9.docx")
