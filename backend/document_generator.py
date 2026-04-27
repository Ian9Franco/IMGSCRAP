import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_property_doc(path, title, description, details_list=None):
    """
    Genera el documento Word de la propiedad interpretando el Markdown básico 
    (##, ###, **, ---) para lograr una jerarquía visual óptima.
    """
    try:
        doc = Document()
        
        # Estilo base
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        if not description:
            description = ""

        # Procesar el texto línea por línea
        lines = description.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph() # Espacio vacío
                continue

            # 1. Separador horizontal ---
            if line.startswith('---'):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run('________________________________________________')
                run.font.color.rgb = None # Gris o negro
                continue

            # 2. Encabezados ## o ###
            if line.startswith('## '):
                p = doc.add_heading(line.replace('## ', ''), level=1)
                continue
            if line.startswith('### '):
                p = doc.add_heading(line.replace('### ', ''), level=2)
                continue

            # 3. Párrafos normales con soporte básico de **negrita**
            p = doc.add_paragraph()
            
            # Regex para encontrar texto entre **
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part.replace('**', ''))
                    run.bold = True
                else:
                    run = p.add_run(part)
                
            # Aplicar negrita extra si tiene emojis de sección (📍, ✨, 🔍, 💰, 📲)
            if any(emoji in line for emoji in ["📍", "✨", "🔍", "💰", "📲"]):
                for run in p.runs:
                    run.bold = True

        # Determinamos el nombre del archivo con versionado
        base_name = "copy_propiedad"
        extension = ".docx"
        
        # Si el archivo base no existe, lo creamos
        docx_path = os.path.join(path, f"{base_name}{extension}")
        
        if os.path.exists(docx_path):
            # Si ya existe, buscamos la siguiente versión disponible
            v = 2
            while os.path.exists(os.path.join(path, f"{base_name}_V{v}{extension}")):
                v += 1
            docx_path = os.path.join(path, f"{base_name}_V{v}{extension}")
            
        doc.save(docx_path)

        # Guardo también un .txt plano (siempre sobreescribe el 'último' para carga rápida)
        txt_path = os.path.join(path, "copy_propiedad.txt")
        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(description)

        return True, docx_path
    except Exception as e:
        print(f"Error generando documento: {e}")
        return False, str(e)
